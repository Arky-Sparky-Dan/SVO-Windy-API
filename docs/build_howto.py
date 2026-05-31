#!/usr/bin/env python3
"""
build_howto.py — generates docs/how-to.docx
Run from the repo root: python3 docs/build_howto.py
Requires: pip install python-docx Pillow requests
"""

import io
import os
import sys
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

OUT_DIR   = Path(__file__).parent
REPO_ROOT = OUT_DIR.parent
OUT_FILE  = OUT_DIR / "how-to.docx"

# Brand colours
BLUE  = RGBColor(0x00, 0x78, 0xD4)   # header blue
GREY  = RGBColor(0x24, 0x24, 0x24)   # body text
BOX_BG = RGBColor(0xF0, 0xF4, 0xFA)  # note box background


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = BLUE
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.color.rgb = GREY
        run.font.size = Pt(11)
    return p


def add_step_header(doc: Document, number: int, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"Step {number}   {title}")
    run.bold = True
    run.font.size  = Pt(13)
    run.font.color.rgb = BLUE
    return p


def add_numbered(doc: Document, items: list):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(11)


def add_bullet(doc: Document, items: list):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(11)


def add_code(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p


def add_note(doc: Document, text: str):
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "EEF4FF")
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent  = Inches(0.1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("ℹ  " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x4E, 0xAA)
    doc.add_paragraph()   # spacing after table


def add_image_from_url(doc: Document, url: str, width: Inches, caption: str = ""):
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = resp.read()
        doc.add_picture(io.BytesIO(data), width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cp.runs:
                run.font.size  = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    except Exception as exc:
        doc.add_paragraph(f"[Image placeholder — {caption or url}]")
        print(f"  Warning: could not fetch image {url}: {exc}")


def add_image_from_file(doc: Document, path: str, width: Inches, caption: str = ""):
    try:
        doc.add_picture(path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cp.runs:
                run.font.size   = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    except Exception as exc:
        doc.add_paragraph(f"[Image placeholder — {caption}]")
        print(f"  Warning: could not add image {path}: {exc}")


def add_divider(doc: Document):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "CCCCCC")
    pb.append(bot)
    pPr.append(pb)


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

def build():
    print("Building how-to.docx ...")
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Cover / title ──────────────────────────────────────────────────────
    title = doc.add_heading("SVO Observatory", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = BLUE
        run.font.size = Pt(28)

    sub = doc.add_paragraph("Windy.com Webcam — Setup Guide")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # Live allsky image
    add_image_from_file(doc, "/tmp/allsky_current.jpg", Inches(3.5),
                        "Current live view from the ASI 224MC — svo.space/allsky/Current.jpg")
    doc.add_paragraph()
    add_divider(doc)

    # ── Overview ──────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Overview", level=1)
    add_body(doc,
        "This guide walks through connecting the ZWO ASI 224MC all-sky camera "
        "at SVO Observatory to the Windy.com live webcam map. Once set up, your "
        "sky image will appear on Windy's global camera map and update automatically "
        "every 60 seconds."
    )
    doc.add_paragraph()

    add_heading(doc, "How it works", level=2)
    add_body(doc,
        "Windy.com uses a pull model — you provide a stable public URL pointing "
        "to your latest image, and Windy fetches it on their own schedule. "
        "No upload script or API key is required."
    )
    doc.add_paragraph()

    # Flow diagram as a table
    tbl = doc.add_table(rows=5, cols=1)
    tbl.style = "Table Grid"
    labels = [
        ("ASI 224MC Camera", "DBEAFE"),
        ("AllSky software  (captures image & uploads to Hostinger via FTP)", "DBEAFE"),
        ("https://svo.space/allsky/Current.jpg  — public image URL", "D1FAE5"),
        ("Windy.com fetches the image every ~60 seconds", "FEF3C7"),
        ("Your camera appears on windy.com/webcams", "D1FAE5"),
    ]
    for i, (label, color) in enumerate(labels):
        cell = tbl.cell(i, 0)
        set_cell_bg(cell, color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.size = Pt(10)
        run.bold = (i in (0, 2, 4))

    doc.add_paragraph()
    add_note(doc,
        "The Windy API key is NOT needed for webcam registration. "
        "It is only used if you later want to query Windy's webcam database programmatically."
    )

    # ── Requirements ─────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "What you need before you start", level=1)
    add_bullet(doc, [
        "Access to the Hostinger control panel (hPanel) — hpanel.hostinger.com",
        "AllSky running on the observatory PC and confirmed uploading Current.jpg to Hostinger",
        "A Windy.com account for the webcam registration form",
        "Verify the live image loads: https://svo.space/allsky/Current.jpg",
    ])

    # ── Step 1 ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_step_header(doc, 1, "Fix the image cache on Hostinger")
    add_body(doc,
        "By default, Hostinger tells browsers and CDNs to cache Current.jpg for 7 days. "
        "This means Windy would show a week-old image. We fix this by adding a small "
        "configuration file to the allsky folder on your web hosting."
    )

    add_heading(doc, "Open Hostinger File Manager", level=2)
    add_numbered(doc, [
        "Go to hpanel.hostinger.com and log in.",
        "Click Hosting in the left menu, then select your hosting plan.",
        'Click "File Manager" to open it.',
    ])
    add_image_from_url(doc,
        "https://www.hostinger.com/tutorials/wp-content/uploads/sites/2/2022/03/hostinger-hpanel-file-manager.png",
        Inches(4.5), "Hostinger hPanel — File Manager location")

    doc.add_paragraph()
    add_heading(doc, "Navigate to the allsky folder", level=2)
    add_numbered(doc, [
        'In the File Manager, navigate to public_html → allsky.',
        "You should see Current.jpg listed there — this is the live image AllSky uploads.",
    ])

    doc.add_paragraph()
    add_heading(doc, "Create the .htaccess file", level=2)
    add_numbered(doc, [
        'Click "New File" at the top of the File Manager.',
        'Name the file exactly:  .htaccess  (with the dot at the start)',
        "Paste the following content into the file editor:",
    ])
    add_code(doc,
        "<Files \"Current.jpg\">\n"
        "    Header set Cache-Control \"public, max-age=30, must-revalidate\"\n"
        "    Header unset Pragma\n"
        "</Files>"
    )
    add_numbered(doc, ["Click Save."])

    doc.add_paragraph()
    add_heading(doc, "Verify it worked", level=2)
    add_body(doc, "Open a browser or command prompt and run:")
    add_code(doc, "curl -sI https://svo.space/allsky/Current.jpg | grep cache-control")
    add_body(doc, "You should see:  cache-control: public, max-age=30, must-revalidate")
    add_note(doc,
        "The .htaccess file is also included in this repo at hostinger/.htaccess "
        "if you need to reference it."
    )

    # ── Step 2 ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_step_header(doc, 2, "Register your camera on Windy.com")
    add_body(doc,
        "This is a one-time step. You fill in a short form on Windy's website, "
        "submit it, and wait 1–3 days for their team to approve it."
    )

    add_numbered(doc, [
        "Go to:  windy.com/webcams/add",
        "Log in to your Windy.com account if prompted.",
        "Fill in the form with the details below.",
        "Click Submit.",
    ])

    doc.add_paragraph()
    add_heading(doc, "Form details to enter", level=2)

    tbl2 = doc.add_table(rows=5, cols=2)
    tbl2.style = "Table Grid"
    headers = [("Field", "Value to enter"),
               ("Location", "SVO Observatory GPS coordinates"),
               ("Camera name", "SVO Observatory All-Sky Camera"),
               ("Description", "ZWO ASI 224MC 180° all-sky camera"),
               ("Image URL", "https://svo.space/allsky/Current.jpg")]
    for i, (field, value) in enumerate(headers):
        row = tbl2.rows[i]
        set_cell_bg(row.cells[0], "DBEAFE" if i == 0 else "F8FAFF")
        set_cell_bg(row.cells[1], "DBEAFE" if i == 0 else "FFFFFF")
        for j, text in enumerate([field, value]):
            p = row.cells[j].paragraphs[0]
            run = p.add_run(text)
            run.bold = (i == 0)
            run.font.size = Pt(10)

    doc.add_paragraph()
    add_note(doc,
        "After submitting, Windy will email you when the webcam is approved. "
        "Once approved your camera appears automatically on windy.com/webcams — "
        "no further configuration needed."
    )

    # ── Step 3 ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_step_header(doc, 3, "Optional — Run the image freshness monitor")
    add_body(doc,
        "The allsky_windy.py script is an optional tool that watches the AllSky "
        "output file on your observatory PC and logs a warning if the image stops "
        "updating. Windy does NOT require this to run — it is purely for your own "
        "peace of mind."
    )

    add_heading(doc, "Install Python (if not already installed)", level=2)
    add_numbered(doc, [
        "Download Python 3.9 or newer from python.org/downloads",
        'During install, tick "Add Python to PATH".',
        "Verify in a Command Prompt:  python --version",
    ])

    doc.add_paragraph()
    add_heading(doc, "Set up the script", level=2)
    add_numbered(doc, [
        "Download the repo ZIP from GitHub (green Code button → Download ZIP).",
        "Extract to  C:\\allsky-windy\\",
        "Copy config.ini.example to config.ini",
        "Open config.ini in Notepad and set image_path to where AllSky writes its image.",
        "Test it by opening a Command Prompt and running:",
    ])
    add_code(doc, r"cd C:\allsky-windy" + "\npython allsky_windy.py")
    add_body(doc, "You should see a line like:")
    add_code(doc, "2026-06-01 10:00:01  INFO      Image updated — 208220 bytes  (update #1)")
    add_numbered(doc, ["Press Ctrl+C to stop. If it works, proceed to auto-start below."])

    doc.add_paragraph()
    add_heading(doc, "Auto-start on boot (Windows Task Scheduler)", level=2)
    add_numbered(doc, [
        "Open the windows\\ folder inside the extracted repo.",
        "Edit allsky-windy-task.xml — update the Python path and script path to match your installation.",
        "Right-click install.bat → Run as administrator.",
        "The script will now start automatically every time the PC boots.",
    ])
    add_code(doc, "schtasks /query /tn \"AllSky-Windy-Uploader\"")
    add_body(doc, "Run the above in Command Prompt to confirm the task is registered.")

    # ── Troubleshooting ───────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Troubleshooting", level=1)

    problems = [
        ("Windy shows a stale or old image",
         "The .htaccess cache fix may not be in place. "
         "Run:  curl -sI https://svo.space/allsky/Current.jpg | grep cache-control\n"
         "It should say max-age=30. If it still says 604800, re-check the .htaccess file in Hostinger."),
        ("Current.jpg on svo.space is not updating",
         "Check AllSky is running on the observatory PC and its FTP upload is configured. "
         "The monitor script will log a warning after 120 seconds of no update."),
        ("Windy registration not approved after 3+ days",
         "Email webcams@windy.com referencing the URL you submitted."),
        ("Monitor script says 'Image not found'",
         r"Check the image_path in config.ini. "
         r"Use PowerShell to search:  Get-ChildItem C:\ -Recurse -Filter image.jpg"),
        ("Task Scheduler task won't start",
         "Re-run install.bat as Administrator. "
         "Open Task Scheduler and check the Last Run Result for AllSky-Windy-Uploader."),
    ]

    for problem, solution in problems:
        p = doc.add_paragraph()
        run = p.add_run(problem)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        add_body(doc, solution)
        doc.add_paragraph()

    # ── Issue reporting ────────────────────────────────────────────────────
    add_divider(doc)
    add_heading(doc, "Reporting issues", level=1)
    add_body(doc,
        "If something isn't working, drop a screenshot in the screenshots/new/ folder "
        "and any log output or notes in context/new/, then open a pull request on GitHub. "
        "Dan and the team will diagnose and push a fix."
    )
    add_code(doc,
        "screenshots/new/   ← screenshot images\n"
        "context/new/       ← log output, notes, error messages"
    )

    # ── Footer note ───────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph("SVO Observatory — github.com/Arky-Sparky-Dan/SVO-Windy-API")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(str(OUT_FILE))
    print(f"Saved: {OUT_FILE}")


if __name__ == "__main__":
    build()
